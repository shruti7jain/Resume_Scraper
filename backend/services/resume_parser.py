"""
Resume Parser Service — Ingestion Layer
Phase 2 Implementation

Handles PDF and DOCX file parsing and returns a structured Resume object.
Uses pdfplumber for PDFs and python-docx for DOCX files.
Section detection is regex/heuristic-based.
"""

import io
import re
from typing import List, Tuple

import pdfplumber
from docx import Document

from backend.models.resume import Resume, ExperienceEntry, EducationEntry, ProjectEntry

# ---------------------------------------------------------------------------
# Section heading patterns
# ---------------------------------------------------------------------------
SECTION_PATTERNS = {
    "summary":        re.compile(r"^\s*(professional\s+summary|summary|objective|profile|about\s+me|executive\s+summary)\s*:?\s*$", re.I),
    "experience":     re.compile(r"^\s*(work\s+experience|experience|employment|work\s+history|career|professional\s+experience)\s*:?\s*$", re.I),
    "skills":         re.compile(r"^\s*(skills|technical\s+skills|core\s+competencies|competencies|expertise|technical\s+expertise|technologies|tech\s+stack|areas\s+of\s+expertise|it\s+skills|professional\s+skills)\s*:?\s*$", re.I),
    "education":      re.compile(r"^\s*(education|academic|qualifications?|schooling|academic\s+background)\s*:?\s*$", re.I),
    "projects":       re.compile(r"^\s*(projects?|personal\s+projects?|key\s+projects?|portfolio|academic\s+projects?)\s*:?\s*$", re.I),
    "certifications": re.compile(r"^\s*(certifications?|certificates?|courses?|achievements?|awards?|licenses?)\s*:?\s*$", re.I),
}

# Patterns for contact extraction
EMAIL_RE   = re.compile(r"[\w.\-+]+@[\w.\-]+\.[a-zA-Z]{2,}")
PHONE_RE   = re.compile(r"(\+?\d[\d\s\-().]{7,}\d)")
LINKEDIN_RE = re.compile(r"linkedin\.com/in/[\w\-]+", re.I)
GITHUB_RE  = re.compile(r"github\.com/[\w\-]+", re.I)
NAME_RE    = re.compile(r"^[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3}$")

# Duration patterns inside experience entries
DURATION_RE = re.compile(
    r"(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*\d{4}"
    r"\s*[–\-—]\s*"
    r"(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*\d{4}"
    r"|"
    r"(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*\d{4}"
    r"\s*[–\-—]\s*present",
    re.I,
)

YEAR_RANGE_RE = re.compile(r"\b(19|20)\d{2}\s*[–\-—]\s*(19|20)\d{2}|\b(19|20)\d{2}\s*[–\-—]\s*present", re.I)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse(file_bytes: bytes, content_type: str) -> Resume:
    """
    Parse a resume file (PDF or DOCX) and return a structured Resume.

    Args:
        file_bytes: Raw file content as bytes
        content_type: MIME type string ('application/pdf' or 'application/vnd.openxmlformats-...')

    Returns:
        Resume: Populated Pydantic model with extracted fields
    """
    if "pdf" in content_type:
        raw_text = _extract_pdf_text(file_bytes)
    else:
        raw_text = _extract_docx_text(file_bytes)

    return _parse_text(raw_text)


# ---------------------------------------------------------------------------
# Extraction helpers
# ---------------------------------------------------------------------------

def _extract_pdf_text(file_bytes: bytes) -> str:
    """Extract raw text from a PDF file using pdfplumber."""
    lines = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text(x_tolerance=2, y_tolerance=2)
            if text:
                lines.append(text)
    return "\n".join(lines)


def _extract_docx_text(file_bytes: bytes) -> str:
    """Extract raw text from a DOCX file using python-docx."""
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                paragraphs.append("  ".join(row_texts))
    return "\n".join(paragraphs)


# ---------------------------------------------------------------------------
# Parsing logic
# ---------------------------------------------------------------------------

def _parse_text(raw_text: str) -> Resume:
    """Parse raw resume text into a structured Resume object."""
    lines = [line.rstrip() for line in raw_text.splitlines()]

    # --- Contact Info (from first 15 lines) ---
    header_block = "\n".join(lines[:15])
    name    = _extract_name(lines[:5])
    email   = _first_match(EMAIL_RE, header_block)
    phone   = _first_match(PHONE_RE, header_block)
    linkedin = _first_match(LINKEDIN_RE, header_block)
    github   = _first_match(GITHUB_RE, header_block)

    # --- Segment into sections ---
    sections = _segment_sections(lines)

    summary        = _parse_summary(sections.get("summary", []))
    experience     = _parse_experience(sections.get("experience", []))
    skills         = _parse_skills(sections.get("skills", []))
    education      = _parse_education(sections.get("education", []))
    projects       = _parse_projects(sections.get("projects", []))
    certifications = _parse_certifications(sections.get("certifications", []))

    return Resume(
        name=name,
        email=email,
        phone=phone,
        linkedin=f"https://{linkedin}" if linkedin and not linkedin.startswith("http") else linkedin,
        github=f"https://{github}" if github and not github.startswith("http") else github,
        professional_summary=summary,
        experience=experience,
        skills=skills,
        education=education,
        projects=projects,
        certifications=certifications,
    )


def _segment_sections(lines: List[str]) -> dict:
    """Split lines into named sections based on heading detection."""
    sections: dict = {}
    current_section = "header"
    sections[current_section] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            sections[current_section].append(line)
            continue

        matched = None
        for section_name, pattern in SECTION_PATTERNS.items():
            if pattern.match(stripped):
                matched = section_name
                break

        if matched:
            current_section = matched
            sections[current_section] = []
        else:
            sections[current_section].append(line)

    return sections


def _extract_name(first_lines: List[str]) -> str | None:
    """Attempt to extract the candidate's name from the top of the resume."""
    for line in first_lines:
        stripped = line.strip()
        if stripped and NAME_RE.match(stripped) and len(stripped.split()) >= 2:
            return stripped
    # Fallback: take the first non-empty, non-contact line
    for line in first_lines:
        stripped = line.strip()
        if stripped and not EMAIL_RE.search(stripped) and not PHONE_RE.search(stripped):
            if len(stripped) < 60:
                return stripped
    return None


def _first_match(pattern: re.Pattern, text: str) -> str | None:
    m = pattern.search(text)
    return m.group(0) if m else None


def _parse_summary(lines: List[str]) -> str:
    """Join summary section lines into a single string."""
    text = " ".join(l.strip() for l in lines if l.strip())
    return text


def _parse_skills(lines: List[str]) -> List[str]:
    """Extract skills from comma/pipe/bullet-separated lines."""
    skills = []
    for line in lines:
        stripped = line.strip().lstrip("•·-–—▪▸►")
        if not stripped:
            continue
        # Split on common delimiters
        parts = re.split(r"[,|;•·▪]", stripped)
        for part in parts:
            skill = part.strip()
            if skill and len(skill) < 60:
                skills.append(skill)
    # Deduplicate preserving order
    seen = set()
    unique = []
    for s in skills:
        key = s.lower()
        if key not in seen:
            seen.add(key)
            unique.append(s)
    return unique


def _parse_experience(lines: List[str]) -> List[ExperienceEntry]:
    """
    Parse experience section into structured ExperienceEntry objects.
    Heuristic: a new entry starts when we detect a company/role line
    (no bullet prefix), followed by a duration pattern.
    """
    entries: List[ExperienceEntry] = []
    current: dict | None = None
    bullets: List[str] = []

    def flush():
        if current:
            entries.append(ExperienceEntry(
                company=current.get("company", "Unknown"),
                role=current.get("role", ""),
                duration=current.get("duration", ""),
                location=current.get("location"),
                bullets=bullets[:],
            ))

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Detect bullet points
        if re.match(r"^[•·▪▸►\-–—*]\s+", stripped):
            bullet_text = re.sub(r"^[•·▪▸►\-–—*]\s+", "", stripped)
            if current:
                bullets.append(bullet_text)
            continue

        # Detect duration in line
        duration_match = DURATION_RE.search(stripped) or YEAR_RANGE_RE.search(stripped)

        if duration_match:
            # Save previous entry
            flush()
            bullets = []
            duration_str = duration_match.group(0).strip()
            # Remove duration from line to get company/role
            remainder = stripped[:duration_match.start()].strip(" ,|–—-")
            # Try to split remainder into role @ company
            if " at " in remainder.lower():
                parts = re.split(r"\s+at\s+", remainder, flags=re.I, maxsplit=1)
                role, company = parts[0].strip(), parts[1].strip()
            elif " | " in remainder or " – " in remainder or " - " in remainder:
                parts = re.split(r"\s*[|–\-]\s*", remainder, maxsplit=1)
                role, company = parts[0].strip(), parts[1].strip() if len(parts) > 1 else ""
            else:
                role, company = remainder, ""
            current = {"role": role, "company": company, "duration": duration_str}
        else:
            # Could be a continuation line (company name, location, etc.)
            if current and not bullets:
                # Likely location or additional info
                if not current.get("company") or not current["company"]:
                    current["company"] = stripped
                elif len(stripped) < 60:
                    current["location"] = stripped
            elif not current:
                # Role/company header before duration line
                current = {"role": stripped, "company": "", "duration": ""}

    flush()
    return entries


def _parse_education(lines: List[str]) -> List[EducationEntry]:
    """Parse education section into EducationEntry objects."""
    entries: List[EducationEntry] = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        year_match = YEAR_RANGE_RE.search(line) or re.search(r"\b(19|20)\d{2}\b", line)
        year_str = year_match.group(0) if year_match else ""

        # Check if this line or the next contains degree keywords
        degree_keywords = re.compile(r"\b(b\.?tech|b\.?sc|b\.?e\b|m\.?tech|m\.?sc|mba|phd|bachelor|master|diploma|associate)\b", re.I)

        if degree_keywords.search(line):
            degree = line
            institution = lines[i + 1].strip() if i + 1 < len(lines) else ""
            i += 2
        elif i + 1 < len(lines) and degree_keywords.search(lines[i + 1]):
            institution = line
            degree = lines[i + 1].strip()
            i += 2
        else:
            # Generic: treat current line as institution
            institution = line
            degree = lines[i + 1].strip() if i + 1 < len(lines) else ""
            i += 2

        if institution or degree:
            entries.append(EducationEntry(
                institution=institution,
                degree=degree,
                year=year_str,
            ))

    return entries


def _parse_projects(lines: List[str]) -> List[ProjectEntry]:
    """Parse projects section."""
    entries: List[ProjectEntry] = []
    current_name = None
    current_desc_lines: List[str] = []
    tech_stack: List[str] = []

    def flush():
        if current_name:
            desc = " ".join(current_desc_lines)
            entries.append(ProjectEntry(
                name=current_name,
                description=desc,
                tech_stack=tech_stack[:],
            ))

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        is_bullet = bool(re.match(r"^[•·▪▸►\-–—*]\s+", stripped))
        if is_bullet:
            bullet_text = re.sub(r"^[•·▪▸►\-–—*]\s+", "", stripped)
            # Check if it's a tech stack line
            if re.search(r"\btech(nologies|nology|nical)?\b.*:|^stack:|^tools?:", bullet_text, re.I):
                tech_part = re.split(r":\s*", bullet_text, maxsplit=1)
                if len(tech_part) > 1:
                    tech_stack = [t.strip() for t in re.split(r"[,|]", tech_part[1]) if t.strip()]
            else:
                current_desc_lines.append(bullet_text)
        else:
            # New project heading
            flush()
            current_name = stripped
            current_desc_lines = []
            tech_stack = []

    flush()
    return entries


def _parse_certifications(lines: List[str]) -> List[str]:
    """Extract certifications as a list of strings."""
    certs = []
    for line in lines:
        stripped = line.strip().lstrip("•·-–—▪▸►*")
        if stripped and len(stripped) > 3:
            certs.append(stripped.strip())
    return certs
